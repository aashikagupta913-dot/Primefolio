import io
from app.core.exceptions import ParsingException
from pypdf import PdfReader
import docx

class ParserService:
    def extract_text(self, file_bytes: bytes, file_name: str) -> str:
        """
        Diverges parser type based on file extension and extracts raw text
        """
        fn_lower = file_name.lower()
        if fn_lower.endswith(".pdf"):
            return self._extract_from_pdf(file_bytes)
        elif fn_lower.endswith(".docx"):
            return self._extract_from_docx(file_bytes)
        elif fn_lower.endswith(".txt") or fn_lower.endswith(".md"):
            return self._extract_from_text(file_bytes)
        else:
            raise ParsingException(
                f"Unsupported file type '{file_name}'. Only .pdf, .docx, and .txt files are supported."
            )

    def _extract_from_pdf(self, file_bytes: bytes) -> str:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            raw_text = "\n".join(text_parts).strip()
            if not raw_text:
                raise ParsingException("PDF file appears to be empty or scanned image only.")
            return raw_text
        except ParsingException:
            raise
        except Exception as e:
            raise ParsingException(f"Error parsing PDF document: {str(e)}")

    def _extract_from_docx(self, file_bytes: bytes) -> str:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_parts = [p.text for p in doc.paragraphs]
            
            # Extract from tables too if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)
                            
            raw_text = "\n".join(text_parts).strip()
            if not raw_text:
                raise ParsingException("Word document (.docx) appears to be empty.")
            return raw_text
        except ParsingException:
            raise
        except Exception as e:
            raise ParsingException(f"Error parsing Word document: {str(e)}")

    def _extract_from_text(self, file_bytes: bytes) -> str:
        try:
            return file_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                return file_bytes.decode("latin-1").strip()
            except Exception as e:
                raise ParsingException(f"Failed to decode text file: {str(e)}")
        except Exception as e:
            raise ParsingException(f"Error reading text file: {str(e)}")

# Global Singleton
parser_service = ParserService()
